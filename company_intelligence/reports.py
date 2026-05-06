from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from company_intelligence.data import FACTOR_DEFINITIONS, format_currency


def factor_weight_summary(weights: dict[str, int]) -> str:
    return ", ".join(
        f"{factor['label']}: {weights.get(factor['key'], 0)}"
        for factor in FACTOR_DEFINITIONS
    )


def build_360_prompt(
    record: Any,
    weights: dict[str, int],
    web_context: str,
    analyst_context: str,
    support_blob: str,
) -> str:
    return f"""
Prepare a professional 360 Degree Analysis for GreenFlow Ventures.

COMPANY DOSSIER
- Company name: {record["company_name"]}
- CIN / LLPIN: {record["cin"]}
- Entity type: {record["entity_type"]}
- State: {record["state"]}
- Sector: {record["sector"]}
- Status: {record["status"]}
- Paid up capital: {format_currency(record["paid_up_capital"])}
- Authorized capital / contribution: {format_currency(record["authorized_capital"])}
- Activity description: {record["activity_description"]}
- Registered address: {record["registered_address"]}
- Email: {record["email"] or "Not available"}
- GreenFlow score: {record["score"]:.1f}/100
- Factor weights: {factor_weight_summary(weights)}

PUBLIC RESEARCH CONTEXT
{web_context}

ANALYST NOTES
{analyst_context or "None provided."}

OPTIONAL SUPPORTING MATERIAL
{support_blob[:6000] if support_blob else "None provided."}

Write in sharp, senior-investment-banking language. Use markdown headings. Cover:
1. Executive summary
2. Business model and value proposition
3. Market positioning and defensibility
4. Promoter / director / management observations
5. Risk register
6. Compliance and IPO readiness signals
7. GreenFlow recommendation and next steps

If something is not certain, label it as an inference. Keep the tone authoritative and commercial.
""".strip()


def build_sector_prompt(record: Any, sector_name: str, web_context: str, focus: str) -> str:
    return f"""
Prepare a sectoral analysis for GreenFlow Ventures.

COMPANY ANCHOR
- Company name: {record["company_name"]}
- Sector selected: {sector_name}
- Activity description: {record["activity_description"]}
- Score: {record["score"]:.1f}/100

PUBLIC RESEARCH CONTEXT
{web_context}

FOCUS AREAS
{focus or "Full sector view."}

Write a professional sector report with markdown headings. Cover:
1. Sector definition and where this company fits
2. Market size and growth signals
3. Demand drivers and policy tailwinds
4. Listed and unlisted peer set
5. Competitive intensity and margin structure
6. Key risks
7. GreenFlow relevance for origination and advisory

State clearly when you are making an inference.
""".strip()


def build_cdr_prompt(record: Any, weights: dict[str, int], web_context: str, focus_areas: str) -> str:
    return f"""
Prepare a Company Detailed Report (CDR) for GreenFlow Ventures.

COMPANY DOSSIER
- Company name: {record["company_name"]}
- CIN / LLPIN: {record["cin"]}
- Entity type: {record["entity_type"]}
- State: {record["state"]}
- Sector: {record["sector"]}
- Status: {record["status"]}
- Paid up capital: {format_currency(record["paid_up_capital"])}
- Authorized capital / contribution: {format_currency(record["authorized_capital"])}
- Activity description: {record["activity_description"]}
- Registered address: {record["registered_address"]}
- Email: {record["email"] or "Not available"}
- GreenFlow score: {record["score"]:.1f}/100
- Factor weights: {factor_weight_summary(weights)}

PUBLIC RESEARCH CONTEXT
{web_context}

PRIORITY FOCUS
{focus_areas or "Full coverage"}

Write a board-quality CDR in markdown with strong headings and subheadings. Cover:
1. Company snapshot
2. Business overview
3. Sector and market context
4. Financial and capital profile
5. Management, promoters, and governance
6. Competitive landscape
7. SWOT
8. Risk register
9. ESG and governance observations
10. IPO readiness and GreenFlow fit
11. Deal thesis
12. Suggested next steps

Use an investment-banking tone. Mark assumptions or inferences clearly.
""".strip()


def _looks_like_heading(line: str) -> int | None:
    stripped = line.strip()
    if not stripped:
        return None
    if stripped.startswith("# "):
        return 1
    if stripped.startswith("## "):
        return 2
    if stripped.startswith("### "):
        return 3
    if re.match(r"^\d+\.\s+[A-Z]", stripped):
        return 2
    if len(stripped) < 90 and stripped.endswith(":"):
        return 3
    return None


def build_docx_report(
    title: str,
    subtitle: str,
    content: str,
    company_snapshot: dict[str, str],
) -> bytes:
    document = Document()
    document.core_properties.title = title
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    document.add_paragraph("")
    document.add_heading("Company Snapshot", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in company_snapshot.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    document.add_paragraph("")
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading_level = _looks_like_heading(line)
        if heading_level:
            clean_heading = re.sub(r"^#{1,3}\s*", "", line).rstrip(":")
            document.add_heading(clean_heading, level=heading_level)
            continue
        if line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue
        document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
