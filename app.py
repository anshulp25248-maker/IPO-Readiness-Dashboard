from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any
from urllib import parse as urlparse

import pandas as pd
import streamlit as st
import requests

try:
    from docx import Document
except Exception:
    Document = None

# ── PAGE CONFIG ──────────────────────────────────────────────
st.set_page_config(
    page_title="GreenFlow Ventures — Company Intelligence",
    page_icon="🌿",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CONSTANTS ────────────────────────────────────────────────
CLAUDE_MODEL  = "claude-sonnet-4-6"
ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"

# ── COMPREHENSIVE INDIA LOCATION DETECTION ──────────────────
# Maps every major city, district, ROC code, and CIN state code
# to the correct Indian state

ROC_TO_STATE = {
    "ROC - AHMEDABAD": "Gujarat", "ROC - ANDAMAN": "Andaman & Nicobar",
    "ROC - BANGALORE": "Karnataka", "ROC - CHANDIGARH": "Punjab",
    "ROC - CHENNAI": "Tamil Nadu", "ROC - CHHATTISGARH": "Chhattisgarh",
    "ROC - COIMBATORE": "Tamil Nadu", "ROC - CUTTACK": "Odisha",
    "ROC - DELHI": "Delhi", "ROC - ERNAKULAM": "Kerala", "ROC - GOA": "Goa",
    "ROC - GWALIOR": "Madhya Pradesh", "ROC - HIMACHAL PRADESH": "Himachal Pradesh",
    "ROC - HYDERABAD": "Telangana", "ROC - JAIPUR": "Rajasthan",
    "ROC - JAMMU": "Jammu & Kashmir", "ROC - JHARKHAND": "Jharkhand",
    "ROC - KANPUR": "Uttar Pradesh", "ROC - KOLKATA": "West Bengal",
    "ROC - MUMBAI": "Maharashtra", "ROC - PATNA": "Bihar",
    "ROC - PONDICHERRY": "Puducherry", "ROC - PUNE": "Maharashtra",
    "ROC - SHILLONG": "Meghalaya", "ROC - UTTARAKHAND": "Uttarakhand",
    "ROC - SHIMLA": "Himachal Pradesh", "ROC - JAMMU & KASHMIR": "Jammu & Kashmir",
    "ROC - LEH": "Ladakh", "ROC - SRINAGAR": "Jammu & Kashmir",
}

CIN_STATE_CODE = {
    "AN": "Andaman & Nicobar", "AP": "Andhra Pradesh", "AR": "Arunachal Pradesh",
    "AS": "Assam", "BR": "Bihar", "CG": "Chhattisgarh", "CH": "Chandigarh",
    "DD": "Daman & Diu", "DL": "Delhi", "DN": "Dadra & Nagar Haveli",
    "GA": "Goa", "GJ": "Gujarat", "HP": "Himachal Pradesh",
    "HR": "Haryana", "JH": "Jharkhand", "JK": "Jammu & Kashmir",
    "KA": "Karnataka", "KL": "Kerala", "LA": "Ladakh", "LD": "Lakshadweep",
    "MH": "Maharashtra", "ML": "Meghalaya", "MN": "Manipur", "MP": "Madhya Pradesh",
    "MZ": "Mizoram", "NL": "Nagaland", "OR": "Odisha", "PB": "Punjab",
    "PY": "Puducherry", "RJ": "Rajasthan", "SK": "Sikkim", "TG": "Telangana",
    "TN": "Tamil Nadu", "TR": "Tripura", "TS": "Telangana", "UK": "Uttarakhand",
    "UP": "Uttar Pradesh", "WB": "West Bengal",
}

# City/district → state mapping (1000+ entries covering all major cities and districts)
CITY_TO_STATE = {
    # Andhra Pradesh
    "visakhapatnam":"Andhra Pradesh","vizag":"Andhra Pradesh","vijayawada":"Andhra Pradesh",
    "guntur":"Andhra Pradesh","nellore":"Andhra Pradesh","kurnool":"Andhra Pradesh",
    "rajahmundry":"Andhra Pradesh","tirupati":"Andhra Pradesh","kakinada":"Andhra Pradesh",
    "anantapur":"Andhra Pradesh","kadapa":"Andhra Pradesh","eluru":"Andhra Pradesh",
    "ongole":"Andhra Pradesh","chittoor":"Andhra Pradesh","srikakulam":"Andhra Pradesh",
    "vizianagaram":"Andhra Pradesh","west godavari":"Andhra Pradesh","east godavari":"Andhra Pradesh",
    "krishna":"Andhra Pradesh","prakasam":"Andhra Pradesh","nandyal":"Andhra Pradesh",
    # Arunachal Pradesh
    "itanagar":"Arunachal Pradesh","tawang":"Arunachal Pradesh","naharlagun":"Arunachal Pradesh",
    # Assam
    "guwahati":"Assam","dibrugarh":"Assam","jorhat":"Assam","silchar":"Assam",
    "tinsukia":"Assam","nagaon":"Assam","bongaigaon":"Assam","kamrup":"Assam",
    "sonitpur":"Assam","lakhimpur":"Assam","cachar":"Assam","golaghat":"Assam",
    "shillong":"Meghalaya","dispur":"Assam",
    # Bihar
    "patna":"Bihar","gaya":"Bihar","bhagalpur":"Bihar","muzaffarpur":"Bihar",
    "darbhanga":"Bihar","purnia":"Bihar","arrah":"Bihar","begusarai":"Bihar",
    "katihar":"Bihar","munger":"Bihar","samastipur":"Bihar","hajipur":"Bihar",
    "chapra":"Bihar","sitamarhi":"Bihar","bettiah":"Bihar","motihari":"Bihar",
    "siwan":"Bihar","nawada":"Bihar","aurangabad":"Bihar","jehanabad":"Bihar",
    # Chhattisgarh
    "raipur":"Chhattisgarh","bhilai":"Chhattisgarh","durg":"Chhattisgarh",
    "bilaspur":"Chhattisgarh","korba":"Chhattisgarh","rajnandgaon":"Chhattisgarh",
    "jagdalpur":"Chhattisgarh","raigarh":"Chhattisgarh","ambikapur":"Chhattisgarh",
    # Delhi NCR
    "delhi":"Delhi","new delhi":"Delhi","dwarka":"Delhi","rohini":"Delhi",
    "okhla":"Delhi","connaught place":"Delhi","noida":"Uttar Pradesh",
    "greater noida":"Uttar Pradesh","gurgaon":"Haryana","gurugram":"Haryana",
    "faridabad":"Haryana","ghaziabad":"Uttar Pradesh","meerut":"Uttar Pradesh",
    # Goa
    "panaji":"Goa","margao":"Goa","vasco":"Goa","panjim":"Goa","mapusa":"Goa",
    "north goa":"Goa","south goa":"Goa",
    # Gujarat
    "ahmedabad":"Gujarat","surat":"Gujarat","vadodara":"Gujarat","rajkot":"Gujarat",
    "bhavnagar":"Gujarat","jamnagar":"Gujarat","junagadh":"Gujarat","gandhinagar":"Gujarat",
    "ankleshwar":"Gujarat","anand":"Gujarat","navsari":"Gujarat","bharuch":"Gujarat",
    "valsad":"Gujarat","mehsana":"Gujarat","morbi":"Gujarat","surendranagar":"Gujarat",
    "amreli":"Gujarat","porbandar":"Gujarat","kutch":"Gujarat","bhuj":"Gujarat",
    "gidc":"Gujarat","gandhi nagar":"Gujarat","deesa":"Gujarat","godhra":"Gujarat",
    "nadiad":"Gujarat","kheda":"Gujarat","patan":"Gujarat","sabarkantha":"Gujarat",
    "banaskantha":"Gujarat","dahod":"Gujarat","panchmahal":"Gujarat",
    "narmada":"Gujarat","tapi":"Gujarat","dangs":"Gujarat","vapi":"Gujarat",
    # Haryana
    "ambala":"Haryana","hisar":"Haryana","rohtak":"Haryana","panipat":"Haryana",
    "karnal":"Haryana","sonipat":"Haryana","yamunanagar":"Haryana","panchkula":"Haryana",
    "bhiwani":"Haryana","sirsa":"Haryana","fatehabad":"Haryana","jhajjar":"Haryana",
    "rewari":"Haryana","mahendragarh":"Haryana","nuh":"Haryana","palwal":"Haryana",
    "kurukshetra":"Haryana","kaithal":"Haryana","jind":"Haryana",
    # Himachal Pradesh
    "shimla":"Himachal Pradesh","dharamshala":"Himachal Pradesh","solan":"Himachal Pradesh",
    "mandi":"Himachal Pradesh","kangra":"Himachal Pradesh","kullu":"Himachal Pradesh",
    "manali":"Himachal Pradesh","hamirpur":"Himachal Pradesh","una":"Himachal Pradesh",
    "baddi":"Himachal Pradesh","palampur":"Himachal Pradesh",
    # Jammu & Kashmir / Ladakh
    "srinagar":"Jammu & Kashmir","jammu":"Jammu & Kashmir","anantnag":"Jammu & Kashmir",
    "baramulla":"Jammu & Kashmir","sopore":"Jammu & Kashmir","kathua":"Jammu & Kashmir",
    "udhampur":"Jammu & Kashmir","leh":"Ladakh","kargil":"Ladakh",
    # Jharkhand
    "ranchi":"Jharkhand","jamshedpur":"Jharkhand","dhanbad":"Jharkhand",
    "bokaro":"Jharkhand","deoghar":"Jharkhand","hazaribagh":"Jharkhand",
    "giridih":"Jharkhand","ramgarh":"Jharkhand","dumka":"Jharkhand",
    # Karnataka
    "bangalore":"Karnataka","bengaluru":"Karnataka","mysore":"Karnataka","mysuru":"Karnataka",
    "hubli":"Karnataka","dharwad":"Karnataka","mangalore":"Karnataka","mangaluru":"Karnataka",
    "belgaum":"Karnataka","belagavi":"Karnataka","gulbarga":"Karnataka","kalaburagi":"Karnataka",
    "davanagere":"Karnataka","bellary":"Karnataka","ballari":"Karnataka",
    "shimoga":"Karnataka","shivamogga":"Karnataka","tumkur":"Karnataka","tumakuru":"Karnataka",
    "udupi":"Karnataka","bidar":"Karnataka","raichur":"Karnataka","koppal":"Karnataka",
    "gadag":"Karnataka","bagalkot":"Karnataka","vijayapura":"Karnataka","bijapur":"Karnataka",
    "hassan":"Karnataka","chikkamagaluru":"Karnataka","kodagu":"Karnataka",
    "kolar":"Karnataka","chickballapur":"Karnataka","bengaluru rural":"Karnataka",
    "ramnagar":"Karnataka","chamarajanagar":"Karnataka","mandya":"Karnataka",
    "yadgir":"Karnataka","haveri":"Karnataka","koramangala":"Karnataka",
    "whitefield":"Karnataka","electronic city":"Karnataka","hsr layout":"Karnataka",
    "indiranagar":"Karnataka","jayanagar":"Karnataka","rajajinagar":"Karnataka",
    # Kerala
    "thiruvananthapuram":"Kerala","kochi":"Kerala","kozhikode":"Kerala",
    "thrissur":"Kerala","kollam":"Kerala","palakkad":"Kerala","alappuzha":"Kerala",
    "malappuram":"Kerala","kannur":"Kerala","kasaragod":"Kerala","kottayam":"Kerala",
    "idukki":"Kerala","pathanamthitta":"Kerala","wayanad":"Kerala","ernakulam":"Kerala",
    "calicut":"Kerala","cochin":"Kerala","trivandrum":"Kerala","thalassery":"Kerala",
    # Madhya Pradesh
    "bhopal":"Madhya Pradesh","indore":"Madhya Pradesh","jabalpur":"Madhya Pradesh",
    "gwalior":"Madhya Pradesh","ujjain":"Madhya Pradesh","sagar":"Madhya Pradesh",
    "dewas":"Madhya Pradesh","satna":"Madhya Pradesh","ratlam":"Madhya Pradesh",
    "rewa":"Madhya Pradesh","murwara":"Madhya Pradesh","singrauli":"Madhya Pradesh",
    "burhanpur":"Madhya Pradesh","khandwa":"Madhya Pradesh","bhind":"Madhya Pradesh",
    "chhindwara":"Madhya Pradesh","shivpuri":"Madhya Pradesh","vidisha":"Madhya Pradesh",
    "chhatarpur":"Madhya Pradesh","damoh":"Madhya Pradesh","mandsaur":"Madhya Pradesh",
    "khargone":"Madhya Pradesh","neemuch":"Madhya Pradesh","pithampur":"Madhya Pradesh",
    "hoshangabad":"Madhya Pradesh","narmadapuram":"Madhya Pradesh",
    # Maharashtra
    "mumbai":"Maharashtra","pune":"Maharashtra","nagpur":"Maharashtra",
    "nashik":"Maharashtra","aurangabad":"Maharashtra","solapur":"Maharashtra",
    "thane":"Maharashtra","kolhapur":"Maharashtra","amravati":"Maharashtra",
    "nanded":"Maharashtra","sangli":"Maharashtra","malegaon":"Maharashtra",
    "jalgaon":"Maharashtra","akola":"Maharashtra","latur":"Maharashtra",
    "dhule":"Maharashtra","ahmednagar":"Maharashtra","chandrapur":"Maharashtra",
    "parbhani":"Maharashtra","ichalkaranji":"Maharashtra","jalna":"Maharashtra",
    "ambarnath":"Maharashtra","bhiwandi":"Maharashtra","panvel":"Maharashtra",
    "navi mumbai":"Maharashtra","vasai":"Maharashtra","mira road":"Maharashtra",
    "kalyan":"Maharashtra","dombivli":"Maharashtra","ulhasnagar":"Maharashtra",
    "badlapur":"Maharashtra","satara":"Maharashtra","beed":"Maharashtra",
    "osmanabad":"Maharashtra","ratnagiri":"Maharashtra","sindhudurg":"Maharashtra",
    "wardha":"Maharashtra","yavatmal":"Maharashtra","buldhana":"Maharashtra",
    "washim":"Maharashtra","hingoli":"Maharashtra","gondia":"Maharashtra",
    "bhandara":"Maharashtra","gadchiroli":"Maharashtra","midc":"Maharashtra",
    "bkc":"Maharashtra","andheri":"Maharashtra","kurla":"Maharashtra",
    "worli":"Maharashtra","powai":"Maharashtra","vikhroli":"Maharashtra",
    "mulund":"Maharashtra","borivali":"Maharashtra","malad":"Maharashtra",
    "kandivali":"Maharashtra","dahisar":"Maharashtra","goregaon":"Maharashtra",
    "jogeshwari":"Maharashtra","bandra":"Maharashtra","juhu":"Maharashtra",
    "vashi":"Maharashtra","belapur":"Maharashtra","airoli":"Maharashtra",
    "khopoli":"Maharashtra","talegaon":"Maharashtra","chakan":"Maharashtra",
    "hinjewadi":"Maharashtra","pimpri":"Maharashtra","chinchwad":"Maharashtra",
    "hadapsar":"Maharashtra","kharadi":"Maharashtra","wakad":"Maharashtra",
    "magarpatta":"Maharashtra","kothrud":"Maharashtra","shivajinagar":"Maharashtra",
    "jnpt":"Maharashtra","nhava sheva":"Maharashtra","raigad":"Maharashtra",
    # Manipur
    "imphal":"Manipur","thoubal":"Manipur","bishnupur":"Manipur",
    # Meghalaya
    "shillong":"Meghalaya","tura":"Meghalaya","jowai":"Meghalaya",
    # Mizoram
    "aizawl":"Mizoram","lunglei":"Mizoram",
    # Nagaland
    "kohima":"Nagaland","dimapur":"Nagaland",
    # Odisha
    "bhubaneswar":"Odisha","cuttack":"Odisha","rourkela":"Odisha",
    "berhampur":"Odisha","sambalpur":"Odisha","puri":"Odisha",
    "balasore":"Odisha","baripada":"Odisha","bhadrak":"Odisha",
    "jharsuguda":"Odisha","angul":"Odisha","dhenkanal":"Odisha",
    "kendrapara":"Odisha","jajpur":"Odisha","koraput":"Odisha",
    "sundargarh":"Odisha","keonjhar":"Odisha","mayurbhanj":"Odisha",
    # Punjab
    "ludhiana":"Punjab","amritsar":"Punjab","jalandhar":"Punjab",
    "patiala":"Punjab","bathinda":"Punjab","hoshiarpur":"Punjab",
    "mohali":"Punjab","ropar":"Punjab","fatehgarh sahib":"Punjab",
    "moga":"Punjab","ferozepur":"Punjab","muktsar":"Punjab",
    "fazilka":"Punjab","pathankot":"Punjab","gurdaspur":"Punjab",
    "kapurthala":"Punjab","sangrur":"Punjab","mansa":"Punjab",
    "nawanshahr":"Punjab","sahibzada ajit singh nagar":"Punjab",
    # Rajasthan
    "jaipur":"Rajasthan","jodhpur":"Rajasthan","kota":"Rajasthan",
    "bikaner":"Rajasthan","ajmer":"Rajasthan","udaipur":"Rajasthan",
    "bhilwara":"Rajasthan","alwar":"Rajasthan","bharatpur":"Rajasthan",
    "sikar":"Rajasthan","pali":"Rajasthan","tonk":"Rajasthan",
    "sri ganganagar":"Rajasthan","hanumangarh":"Rajasthan","banswara":"Rajasthan",
    "baran":"Rajasthan","barmer":"Rajasthan","bundi":"Rajasthan",
    "chittorgarh":"Rajasthan","churu":"Rajasthan","dausa":"Rajasthan",
    "dholpur":"Rajasthan","dungarpur":"Rajasthan","jaisalmer":"Rajasthan",
    "jalore":"Rajasthan","jhalawar":"Rajasthan","jhunjhunu":"Rajasthan",
    "karauli":"Rajasthan","nagaur":"Rajasthan","pratapgarh":"Rajasthan",
    "rajsamand":"Rajasthan","sawai madhopur":"Rajasthan","sirohi":"Rajasthan",
    "jaipur tech park":"Rajasthan","jaipur sez":"Rajasthan",
    # Sikkim
    "gangtok":"Sikkim","namchi":"Sikkim",
    # Tamil Nadu
    "chennai":"Tamil Nadu","coimbatore":"Tamil Nadu","madurai":"Tamil Nadu",
    "tiruchirappalli":"Tamil Nadu","tiruchy":"Tamil Nadu","trichy":"Tamil Nadu",
    "salem":"Tamil Nadu","tirunelveli":"Tamil Nadu","tiruppur":"Tamil Nadu",
    "vellore":"Tamil Nadu","erode":"Tamil Nadu","thoothukudi":"Tamil Nadu",
    "tuticorin":"Tamil Nadu","dindigul":"Tamil Nadu","thanjavur":"Tamil Nadu",
    "ranipet":"Tamil Nadu","sivakasi":"Tamil Nadu","karur":"Tamil Nadu",
    "hosur":"Tamil Nadu","kancheepuram":"Tamil Nadu","krishnagiri":"Tamil Nadu",
    "namakkal":"Tamil Nadu","cuddalore":"Tamil Nadu","pudukottai":"Tamil Nadu",
    "theni":"Tamil Nadu","nagapattinam":"Tamil Nadu","tiruvarur":"Tamil Nadu",
    "perambalur":"Tamil Nadu","ariyalur":"Tamil Nadu","kallakurichi":"Tamil Nadu",
    "villupuram":"Tamil Nadu","chengalpattu":"Tamil Nadu","tirupattur":"Tamil Nadu",
    "madurai":"Tamil Nadu","virudhunagar":"Tamil Nadu","ramanathapuram":"Tamil Nadu",
    "tiruppur":"Tamil Nadu","chennai auto park":"Tamil Nadu","tirupur export zone":"Tamil Nadu",
    # Telangana
    "hyderabad":"Telangana","warangal":"Telangana","nizamabad":"Telangana",
    "karimnagar":"Telangana","khammam":"Telangana","mahbubnagar":"Telangana",
    "nalgonda":"Telangana","adilabad":"Telangana","medak":"Telangana",
    "rangareddy":"Telangana","secunderabad":"Telangana","cyberabad":"Telangana",
    "hitech city":"Telangana","gachibowli":"Telangana","madhapur":"Telangana",
    "begumpet":"Telangana","banjara hills":"Telangana","jubilee hills":"Telangana",
    "mahabubnagar":"Telangana","sangareddy":"Telangana","siddipet":"Telangana",
    "vikarabad":"Telangana","wanaparthy":"Telangana","yadadri":"Telangana",
    "nagarkurnool":"Telangana","suryapet":"Telangana","jangaon":"Telangana",
    "jayashankar":"Telangana","mancherial":"Telangana","asifabad":"Telangana",
    "nirmal":"Telangana","nizamabad":"Telangana","kamareddy":"Telangana",
    "rajanna sircilla":"Telangana","peddapalli":"Telangana","mulugu":"Telangana",
    # Tripura
    "agartala":"Tripura","dharmanagar":"Tripura",
    # Uttar Pradesh
    "lucknow":"Uttar Pradesh","kanpur":"Uttar Pradesh","agra":"Uttar Pradesh",
    "varanasi":"Uttar Pradesh","prayagraj":"Uttar Pradesh","allahabad":"Uttar Pradesh",
    "meerut":"Uttar Pradesh","bareilly":"Uttar Pradesh","aligarh":"Uttar Pradesh",
    "moradabad":"Uttar Pradesh","saharanpur":"Uttar Pradesh","gorakhpur":"Uttar Pradesh",
    "firozabad":"Uttar Pradesh","jhansi":"Uttar Pradesh","muzaffarnagar":"Uttar Pradesh",
    "mathura":"Uttar Pradesh","bijnor":"Uttar Pradesh","rampur":"Uttar Pradesh",
    "shahjahanpur":"Uttar Pradesh","farrukhabad":"Uttar Pradesh","mau":"Uttar Pradesh",
    "hapur":"Uttar Pradesh","etawah":"Uttar Pradesh","mirzapur":"Uttar Pradesh",
    "bulandshahr":"Uttar Pradesh","sambhal":"Uttar Pradesh","amroha":"Uttar Pradesh",
    "hardoi":"Uttar Pradesh","fatehpur":"Uttar Pradesh","raebareli":"Uttar Pradesh",
    "orai":"Uttar Pradesh","sitapur":"Uttar Pradesh","bahraich":"Uttar Pradesh",
    "modinagar":"Uttar Pradesh","unnao":"Uttar Pradesh","jaunpur":"Uttar Pradesh",
    "lakhimpur":"Uttar Pradesh","hathras":"Uttar Pradesh","banda":"Uttar Pradesh",
    "pilibhit":"Uttar Pradesh","barabanki":"Uttar Pradesh","khurja":"Uttar Pradesh",
    "noida":"Uttar Pradesh","greater noida":"Uttar Pradesh","ghaziabad":"Uttar Pradesh",
    "vrindavan":"Uttar Pradesh","brij":"Uttar Pradesh","chitrakoot":"Uttar Pradesh",
    # Uttarakhand
    "dehradun":"Uttarakhand","haridwar":"Uttarakhand","rishikesh":"Uttarakhand",
    "nainital":"Uttarakhand","roorkee":"Uttarakhand","kashipur":"Uttarakhand",
    "rudrapur":"Uttarakhand","haldwani":"Uttarakhand","pantnagar":"Uttarakhand",
    "pithoragarh":"Uttarakhand","almora":"Uttarakhand","mussoorie":"Uttarakhand",
    "kotdwar":"Uttarakhand","srinagar garhwal":"Uttarakhand","pauri":"Uttarakhand",
    "tehri":"Uttarakhand","champawat":"Uttarakhand","bageshwar":"Uttarakhand",
    # West Bengal
    "kolkata":"West Bengal","howrah":"West Bengal","durgapur":"West Bengal",
    "asansol":"West Bengal","siliguri":"West Bengal","bardhaman":"West Bengal",
    "malda":"West Bengal","baharampur":"West Bengal","habra":"West Bengal",
    "kharagpur":"West Bengal","haldia":"West Bengal","raiganj":"West Bengal",
    "midnapore":"West Bengal","medinipur":"West Bengal","jalpaiguri":"West Bengal",
    "cooch behar":"West Bengal","alipurduar":"West Bengal","purulia":"West Bengal",
    "bankura":"West Bengal","bishnupur":"West Bengal","krishnanagar":"West Bengal",
    "burdwan":"West Bengal","ranaghat":"West Bengal","barrackpore":"West Bengal",
    "dum dum":"West Bengal","salt lake":"West Bengal","new town":"West Bengal",
    "rajarhat":"West Bengal","belgharia":"West Bengal","hooghly":"West Bengal",
    "serampore":"West Bengal","bally":"West Bengal","uttarpara":"West Bengal",
    # Union Territories
    "chandigarh":"Chandigarh","mohali":"Punjab","panchkula":"Haryana",
    "silvassa":"Dadra & Nagar Haveli","daman":"Daman & Diu","diu":"Daman & Diu",
    "kavaratti":"Lakshadweep","pondicherry":"Puducherry","puducherry":"Puducherry",
    "port blair":"Andaman & Nicobar",
}

SECTOR_KEYWORDS = {
    "Technology":         ["software","saas","ai ","artificial intelligence","cloud","data analytics","platform","cyber","iot","digital","information technology","it services","computer"],
    "Healthcare":         ["pharma","pharmaceutical","hospital","healthcare","medical","diagnostic","biotech","clinical","device","drug","health"],
    "Renewable Energy":   ["solar","wind","green energy","renewable","clean tech","ev ","battery","hydro","biofuel","clean energy","power generation"],
    "Financial Services": ["fintech","lending","nbfc","insurance","wealth","payment","banking","microfinance","financial service","capital","investment"],
    "Agriculture":        ["agri","agriculture","food processing","dairy","farm","fisheri","horticulture","poultry","organic","allied activities","seeds","irrigation"],
    "Manufacturing":      ["manufactur","industrial","engineering","fabrication","chemical","automotive","steel","metal","auto component","component","casting","forging"],
    "Logistics":          ["logistic","freight","supply chain","transport","warehouse","shipping","courier","cargo","express","forwarding"],
    "Infrastructure":     ["real estate","construction","infra","property","building","housing","developer","road","bridge","urban"],
    "Education":          ["education","edtech","learning","training","school","college","coaching","skill","university","institute"],
    "Consumer":           ["retail","consumer","brand","fmcg","lifestyle","commerce","fashion","apparel","textile","garment"],
    "Defence":            ["defence","defense","aerospace","semiconductor","space","military","security","ordnance"],
}


def detect_state_from_address(address: str) -> str:
    """Detect Indian state from any address string using city/district matching."""
    if not address:
        return ""
    text = address.lower()
    # Remove punctuation for cleaner matching
    text_clean = re.sub(r"[,\-/\|\.]+", " ", text)

    # Check direct state name mentions first
    state_names = {
        "andhra pradesh","arunachal pradesh","assam","bihar","chhattisgarh",
        "goa","gujarat","haryana","himachal pradesh","jharkhand","karnataka",
        "kerala","madhya pradesh","maharashtra","manipur","meghalaya","mizoram",
        "nagaland","odisha","punjab","rajasthan","sikkim","tamil nadu","telangana",
        "tripura","uttar pradesh","uttarakhand","west bengal","delhi","chandigarh",
        "puducherry","ladakh","jammu kashmir","jammu & kashmir",
    }
    for sn in state_names:
        if sn in text_clean:
            return sn.title().replace("&","&")

    # Check state abbreviations at end of address (e.g. "- MH", "AP", "GJ")
    abbr_match = re.search(r'\b([A-Z]{2})\b[-\s]*\d{6}', address.upper())
    if abbr_match:
        code = abbr_match.group(1)
        if code in CIN_STATE_CODE:
            return CIN_STATE_CODE[code]

    # PIN code prefix → state
    pin_match = re.search(r'\b(\d{6})\b', address)
    if pin_match:
        pin = pin_match.group(1)
        pin_state = pin_prefix_to_state(pin)
        if pin_state:
            return pin_state

    # City/district keyword matching
    words = text_clean.split()
    # Try multi-word city names first (up to 3 words)
    for length in [3, 2, 1]:
        for i in range(len(words) - length + 1):
            phrase = " ".join(words[i:i+length]).strip()
            if phrase in CITY_TO_STATE:
                return CITY_TO_STATE[phrase]

    return ""


def pin_prefix_to_state(pin: str) -> str:
    """Map 6-digit PIN code prefix to Indian state."""
    prefix = pin[:3]
    pin_map = {
        # Maharashtra: 400–445
        **{str(p): "Maharashtra" for p in range(400, 446)},
        # Delhi: 110
        "110": "Delhi",
        # Karnataka: 560–591
        **{str(p): "Karnataka" for p in range(560, 592)},
        # Tamil Nadu: 600–643
        **{str(p): "Tamil Nadu" for p in range(600, 644)},
        # Andhra Pradesh / Telangana: 500–535
        **{str(p): "Telangana" for p in range(500, 510)},
        **{str(p): "Andhra Pradesh" for p in range(510, 536)},
        # Gujarat: 360–396
        **{str(p): "Gujarat" for p in range(360, 397)},
        # Rajasthan: 302–345
        **{str(p): "Rajasthan" for p in range(302, 346)},
        # Uttar Pradesh: 200–285
        **{str(p): "Uttar Pradesh" for p in range(200, 286)},
        # West Bengal: 700–743
        **{str(p): "West Bengal" for p in range(700, 744)},
        # Punjab: 140–160
        **{str(p): "Punjab" for p in range(140, 161)},
        # Haryana: 121–136
        **{str(p): "Haryana" for p in range(121, 137)},
        # Madhya Pradesh: 450–488
        **{str(p): "Madhya Pradesh" for p in range(450, 489)},
        # Bihar: 800–855
        **{str(p): "Bihar" for p in range(800, 856)},
        # Kerala: 670–695
        **{str(p): "Kerala" for p in range(670, 696)},
        # Odisha: 751–770
        **{str(p): "Odisha" for p in range(751, 771)},
        # Jharkhand: 814–835
        **{str(p): "Jharkhand" for p in range(814, 836)},
        # Chhattisgarh: 490–497
        **{str(p): "Chhattisgarh" for p in range(490, 498)},
        # Uttarakhand: 246–263
        **{str(p): "Uttarakhand" for p in range(246, 264)},
        # Himachal Pradesh: 171–177
        **{str(p): "Himachal Pradesh" for p in range(171, 178)},
        # Assam: 781–788
        **{str(p): "Assam" for p in range(781, 789)},
        # Goa: 403
        "403": "Goa",
    }
    return pin_map.get(prefix, "")


def parse_cin(cin: str) -> dict:
    m = re.match(r"^[UL]\d{5}([A-Z]{2})(\d{4})([A-Z]{2,3})\d+", (cin or "").strip().upper())
    if not m:
        return {"state": "", "year": None, "cls": "Private"}
    clss = {
        "PTC": "Private", "PLC": "Public", "OPC": "One Person",
        "NPL": "Section 8", "FTC": "Foreign", "ULL": "Unlimited",
    }
    return {
        "state": CIN_STATE_CODE.get(m.group(1), ""),
        "year":  int(m.group(2)),
        "cls":   clss.get(m.group(3), "Private"),
    }


def detect_state(row: pd.Series, cin_data: dict, roc: str) -> str:
    """Try multiple sources to find the state, in order of reliability."""
    # 1. ROC office (most reliable)
    if roc and roc in ROC_TO_STATE:
        return ROC_TO_STATE[roc]
    # 2. CIN state code
    if cin_data.get("state"):
        return cin_data["state"]
    # 3. Explicit State column
    for col in ["State","STATE","state","Registered State"]:
        v = str(row.get(col,"")).strip()
        if v and v.lower() not in ("nan","none",""):
            # Try to match against known states
            vl = v.lower()
            for sn in CIN_STATE_CODE.values():
                if vl in sn.lower() or sn.lower() in vl:
                    return sn
            return v
    # 4. Parse from address
    addr = str(row.get("REGISTERED OFFICE ADDRESS", row.get("Registered Office Address", ""))).strip()
    state_from_addr = detect_state_from_address(addr)
    if state_from_addr:
        return state_from_addr
    return "Unknown"


# ─── FACTOR DEFINITIONS ─────────────────────────────────────
@dataclass
class Factor:
    key: str
    label: str
    weight: float
    desc: str

# 15 factors × 6.66 default weight ≈ 100 total | each slider: 0–10
_W = 10.00
FACTORS = [
    Factor("capital",      "Paid-up Capital",        _W, "₹1Cr–₹25Cr is ideal for SME IPO eligibility"),
    Factor("headroom",     "Capital Headroom",        _W, "Higher auth–paid gap supports fresh equity raising"),
    Factor("vintage",      "Company Vintage",         _W, "3–12 years is the optimal growth window for IPO"),
    Factor("structure",    "Corporate Structure",     _W, "Private Ltd preferred for conversion to Public"),
    Factor("sector",       "Sector Momentum",         _W, "Investor appetite and SEBI listing demand by sector"),
    Factor("status",       "Active Status",           _W, "Active MCA status required for IPO eligibility"),
    Factor("geography",    "Geography",               _W, "Proximity to financial centres improves deal access"),
    Factor("disclosure",   "Disclosure Depth",        _W, "Email, address, website completeness score"),
    Factor("signal",       "Data Confidence",         _W, "Overall MCA record data quality"),
    Factor("profitability","Revenue Potential",       _W, "Capital size and sector as proxy for revenue potential"),
    Factor("compliance",   "MCA Compliance Score",   _W, "Completeness and consistency of MCA filing data"),
    Factor("scalability",  "Business Scalability",   _W, "Sector-based assessment of growth scalability potential"),
    Factor("promoter",     "Promoter Credibility",   _W, "Inferred promoter strength from structure and disclosure"),
    Factor("ipo_readiness","IPO Readiness",           _W, "Composite readiness signal: capital, vintage, structure"),
    Factor("regulatory",   "Regulatory Environment", _W, "Regulatory tailwinds and SEBI listing appetite for sector"),
]

# ─── CSS — GREENFLOW PROFESSIONAL LIGHT THEME ───────────────
def inject_css() -> None:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Nunito+Sans:ital,opsz,wght@0,6..12,300;0,6..12,400;0,6..12,600;0,6..12,700;0,6..12,800&family=Playfair+Display:wght@600;700&family=JetBrains+Mono:wght@400;500;600&display=swap');

    /* ── ROOT PALETTE ── */
    :root {
        --gf-green-dark:    #1A4D2E;
        --gf-green-mid:     #2E7D32;
        --gf-green-bright:  #43A047;
        --gf-green-light:   #C8E6C9;
        --gf-green-pale:    #E8F5E9;
        --gf-green-bg:      #F0F7F1;
        --gf-gold:          #B8860B;
        --gf-gold-light:    #F0E68C;
        --gf-white:         #FFFFFF;
        --gf-text-dark:     #0D2818;
        --gf-text-mid:      #2E4A35;
        --gf-text-muted:    #5A7A5E;
        --gf-border:        #A5D6A7;
        --gf-surface:       #FAFCFA;
    }

    /* ── BASE ── */
    html, body,
    .stApp, [data-testid="stAppViewContainer"],
    [data-testid="stMain"], [data-testid="stAppViewBlockContainer"] {
        background-color: var(--gf-green-bg) !important;
        font-family: 'Nunito Sans', sans-serif !important;
        color: var(--gf-text-dark) !important;
    }
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #E8F5E9 0%, #F0F7F1 100%) !important;
        border-right: 2px solid var(--gf-border) !important;
    }
    [data-testid="stHeader"] {
        background-color: rgba(240,247,241,0.96) !important;
        border-bottom: 1px solid var(--gf-border) !important;
    }

    /* ── TYPOGRAPHY ── */
    h1 {
        font-family: 'Playfair Display', serif !important;
        color: var(--gf-green-dark) !important;
        font-size: 2rem !important;
        letter-spacing: -0.5px !important;
    }
    h2 {
        font-family: 'Playfair Display', serif !important;
        color: var(--gf-green-dark) !important;
        font-size: 1.45rem !important;
    }
    h3 {
        font-family: 'Nunito Sans', sans-serif !important;
        color: var(--gf-green-dark) !important;
        font-size: 1.1rem !important;
        font-weight: 700 !important;
    }

    /* ── TABS ── */
    .stTabs [data-baseweb="tab-list"] {
        gap: 6px !important;
        background: transparent !important;
        border-bottom: 2px solid var(--gf-border) !important;
        padding-bottom: 0 !important;
    }
    .stTabs [data-baseweb="tab"] {
        background: var(--gf-white) !important;
        border: 1.5px solid var(--gf-border) !important;
        border-bottom: none !important;
        border-radius: 8px 8px 0 0 !important;
        padding: 8px 20px !important;
        font-family: 'Nunito Sans', sans-serif !important;
        font-size: 13.5px !important;
        font-weight: 600 !important;
        color: var(--gf-text-mid) !important;
        transition: all 0.18s ease !important;
    }
    .stTabs [data-baseweb="tab"]:hover {
        background: var(--gf-green-pale) !important;
        color: var(--gf-green-dark) !important;
        border-color: var(--gf-green-bright) !important;
    }
    /* SELECTED TAB — light green bg, dark text for maximum readability */
    .stTabs [data-baseweb="tab"][aria-selected="true"],
    .stTabs [aria-selected="true"] {
        background: var(--gf-green-bright) !important;
        border-color: var(--gf-green-mid) !important;
        color: #FFFFFF !important;
        font-weight: 800 !important;
        box-shadow: 0 -2px 0 var(--gf-green-bright) inset, 0 2px 8px rgba(46,125,50,0.22) !important;
    }
    /* Force text inside selected tab to be white */
    .stTabs [aria-selected="true"] p,
    .stTabs [aria-selected="true"] span,
    .stTabs [aria-selected="true"] div {
        color: #FFFFFF !important;
    }
    .stTabs [data-baseweb="tab-panel"] {
        padding-top: 20px !important;
    }

    /* ── BUTTONS ── */
    .stButton > button {
        background: linear-gradient(135deg, var(--gf-green-mid) 0%, var(--gf-green-dark) 100%) !important;
        color: #FFFFFF !important;
        border: 1.5px solid var(--gf-green-bright) !important;
        border-radius: 8px !important;
        font-family: 'Nunito Sans', sans-serif !important;
        font-weight: 700 !important;
        font-size: 13.5px !important;
        padding: 0.55rem 1.2rem !important;
        letter-spacing: 0.3px !important;
        box-shadow: 0 2px 8px rgba(30,100,50,0.2) !important;
        transition: all 0.16s ease !important;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, var(--gf-green-bright) 0%, var(--gf-green-mid) 100%) !important;
        box-shadow: 0 4px 14px rgba(30,100,50,0.32) !important;
        transform: translateY(-1px) !important;
    }
    .stButton > button:disabled {
        background: #D0E9D4 !important;
        color: #8AA88D !important;
        box-shadow: none !important;
        border-color: var(--gf-border) !important;
        transform: none !important;
    }
    .stDownloadButton > button {
        background: var(--gf-white) !important;
        color: var(--gf-green-dark) !important;
        border: 1.5px solid var(--gf-gold) !important;
        border-radius: 8px !important;
        font-weight: 700 !important;
        font-family: 'Nunito Sans', sans-serif !important;
    }
    .stDownloadButton > button:hover {
        background: var(--gf-gold-light) !important;
    }

    /* ── FORM INPUTS ── */
    .stTextInput input, .stTextArea textarea, .stNumberInput input {
        background: var(--gf-white) !important;
        border: 1.5px solid var(--gf-border) !important;
        border-radius: 7px !important;
        color: var(--gf-text-dark) !important;
        font-family: 'Nunito Sans', sans-serif !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--gf-green-bright) !important;
        box-shadow: 0 0 0 2.5px rgba(67,160,71,0.18) !important;
    }
    div[data-baseweb="select"] > div {
        background: var(--gf-white) !important;
        border-color: var(--gf-border) !important;
        color: var(--gf-text-dark) !important;
    }
    [data-testid="stFileUploader"] section {
        background: var(--gf-green-pale) !important;
        border: 2px dashed var(--gf-green-bright) !important;
        border-radius: 10px !important;
    }

    /* ── METRICS ── */
    div[data-testid="stMetric"] {
        background: var(--gf-white) !important;
        border: 1px solid var(--gf-border) !important;
        border-left: 4px solid var(--gf-green-bright) !important;
        border-radius: 10px !important;
        padding: 14px 16px !important;
        box-shadow: 0 1px 4px rgba(30,100,50,0.08) !important;
    }
    div[data-testid="stMetric"] label {
        color: var(--gf-text-muted) !important;
        font-size: 0.74rem !important;
        font-weight: 800 !important;
        letter-spacing: 0.07em !important;
        text-transform: uppercase !important;
    }
    div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
        color: var(--gf-green-dark) !important;
        font-family: 'JetBrains Mono', monospace !important;
        font-size: 1.6rem !important;
        font-weight: 600 !important;
    }

    /* ── DATAFRAME ── */
    [data-testid="stDataFrame"] {
        border-radius: 10px !important;
        border: 1px solid var(--gf-border) !important;
        overflow: hidden !important;
    }

    /* ── MISC ── */
    details {
        border: 1px solid var(--gf-border) !important;
        border-radius: 8px !important;
        background: var(--gf-surface) !important;
    }
    summary { color: var(--gf-green-dark) !important; font-weight: 700 !important; }
    .stSuccess { border-radius: 8px !important; }
    .stWarning { border-radius: 8px !important; }
    .stInfo    { border-radius: 8px !important; }
    hr { border-color: var(--gf-border) !important; border-width: 1px !important; }
    [data-testid="stSidebar"] h3 { color: var(--gf-green-dark) !important; font-size: 0.9rem !important; }
    .stCaption, [data-testid="stCaption"] { color: var(--gf-text-muted) !important; font-size: 0.82rem !important; }
    div[data-testid="stMarkdownContainer"] p  { color: var(--gf-text-dark) !important; }
    div[data-testid="stMarkdownContainer"] li { color: var(--gf-text-dark) !important; }
    div[data-testid="stMarkdownContainer"] strong { color: var(--gf-green-dark) !important; }
    div[data-testid="stMarkdownContainer"] a  { color: var(--gf-green-mid) !important; }
    </style>
    """, unsafe_allow_html=True)


# ─── HELPERS ────────────────────────────────────────────────
def parse_money(v: Any) -> float:
    if v is None: return 0.0
    try: return float(re.sub(r"[^0-9.\-]","",str(v).strip()))
    except: return 0.0

def inr(v: float) -> str:
    if v >= 1e7:  return f"₹{v/1e7:.2f} Cr"
    if v >= 1e5:  return f"₹{v/1e5:.2f} L"
    return f"₹{v:,.0f}"

def safe_fn(t: str) -> str:
    return re.sub(r'[<>:"/\\|?*]+',"",t).strip()[:60] or "greenflow"

def score_label(s: int) -> str:
    if s >= 85: return "Prime"
    if s >= 70: return "High Potential"
    if s >= 55: return "Qualified"
    if s >= 40: return "Watchlist"
    return "Low Priority"

def classify_sector(act: str) -> str:
    t = (act or "").lower()
    for sec, kws in SECTOR_KEYWORDS.items():
        if any(k in t for k in kws): return sec
    return "Other"

def get_val(row: pd.Series, cols: list, default: str="") -> str:
    for c in cols:
        v = str(row.get(c,"")).strip()
        if v and v.lower() not in ("nan","none",""): return v
    return default

def links(name: str, cin: str) -> dict:
    q = urlparse.quote(name or cin)
    return {
        "Google":   f"https://www.google.com/search?q={q}+company+India",
        "News":     f"https://news.google.com/search?q={q}",
        "Zauba":    f"https://www.zaubacorp.com/search/results/{q}",
        "Tofler":   f"https://www.tofler.in/search?query={q}",
        "MCA":      "https://www.mca.gov.in/content/mca/global/en/mca/master-data/MDS.html",
        "LinkedIn": f"https://www.linkedin.com/search/results/companies/?keywords={q}",
    }


# ─── FILE PARSER ────────────────────────────────────────────
def find_header(df: pd.DataFrame) -> pd.DataFrame:
    df = df.dropna(how="all").reset_index(drop=True)
    if df.empty: return df
    tokens = ["company","cin","activity","roc","capital","address","email","llpin","status"]
    bi, bs = 0, -1
    for i in range(min(len(df),15)):
        row_str = " ".join(str(v).lower() for v in df.iloc[i])
        s = sum(1 for t in tokens if t in row_str)
        if s > bs: bs, bi = s, i
    hdr = [str(v).strip() or f"C{j}" for j,v in enumerate(df.iloc[bi])]
    body = df.iloc[bi+1:].reset_index(drop=True)
    body.columns = hdr
    return body.dropna(how="all").reset_index(drop=True)

def parse_rows(df: pd.DataFrame) -> pd.DataFrame:
    df = find_header(df)
    rows = []
    for idx, row in df.iterrows():
        cin    = get_val(row, ["CIN","cin","Cin","LLPIN"])
        cd     = parse_cin(cin)
        roc    = get_val(row, ["ROC","Roc"])
        act    = get_val(row, ["ACTIVITY DESCRIPTION","Principal Business Activity","Business Activity"], "General")
        name   = get_val(row, ["COMPANY NAME","Company Name","COMPANY_NAME","LIMITED LIABILITY PARTNERSHIP NAME"], "Unknown")
        paid   = parse_money(get_val(row, ["PAIDUP CAPITAL","Paid Up Capital","Paid Up Capital (in Rs.)"]))
        auth   = parse_money(get_val(row, ["AUTHORIZED CAPITAL","Authorised Capital","Authorised Capital (in Rs.)"])) or paid
        inc    = get_val(row, ["DATE OF INCORPORATION","DATE OF REGISTRATION"])
        email  = get_val(row, ["EMAIL","Email","Email ID of the Company"])
        addr   = get_val(row, ["REGISTERED OFFICE ADDRESS","Registered Office Address"])
        state  = detect_state(row, cd, roc)
        rows.append({
            "company_name":       name,
            "cin":                cin or f"UNK{idx}",
            "state":              state,
            "roc":                roc,
            "status":             get_val(row, ["Company Status","Company Status (for efiling)"], "Active"),
            "activity":           act,
            "sector":             classify_sector(act),
            "paid_up_capital":    paid,
            "authorised_capital": auth,
            "company_class":      get_val(row, ["Class of Company","CLASS"], cd["cls"]),
            "email":              email,
            "address":            addr,
            "inc_year":           inc or cd["year"] or "",
        })
    out = pd.DataFrame(rows).drop_duplicates(subset=["cin","company_name"]).reset_index(drop=True)
    out["score"] = 0
    return out

def load_file(f) -> pd.DataFrame:
    name = getattr(f,"name","").lower()
    if name.endswith(".csv"):
        return parse_rows(pd.read_csv(f, header=None))
    wb = pd.ExcelFile(f)
    frames = []
    for sn in wb.sheet_names:
        raw = pd.read_excel(wb, sheet_name=sn, header=None)
        if raw.empty: continue
        p = parse_rows(raw)
        if not p.empty:
            p["sheet"] = sn
            frames.append(p)
    if not frames: return pd.DataFrame()
    return pd.concat(frames,ignore_index=True).drop_duplicates(subset=["cin","company_name"]).reset_index(drop=True)


# ─── SCORING ────────────────────────────────────────────────
def score_one(row: pd.Series, W: dict) -> int:
    w = {k:v for k,v in W.items() if v > 0}
    if not w: return 0

    def cap():
        p = float(row.get("paid_up_capital",0) or 0)
        return 100 if 1e7<=p<=2.5e8 else 75 if 5e6<=p<1e7 else 70 if 2.5e8<p<=6e8 else 45 if p>0 else 15

    def headroom():
        p = float(row.get("paid_up_capital",0) or 0)
        a = float(row.get("authorised_capital",0) or 0)
        if a<=0: return 35
        r = max(a-p,0)/a
        return 100 if r>=0.6 else 78 if r>=0.35 else 55 if r>=0.15 else 30

    def vintage():
        y = row.get("inc_year")
        try: age = datetime.now().year - int(str(y)[:4])
        except: return 45
        return 100 if 4<=age<=12 else 72 if 2<=age<4 else 68 if 13<=age<=20 else 40

    def structure():
        c = str(row.get("company_class","")).lower()
        return 100 if "private" in c else 82 if "public" in c else 60 if "one person" in c else 35

    def sector():
        m = {"Technology":94,"Healthcare":90,"Renewable Energy":100,"Financial Services":86,
             "Manufacturing":76,"Agriculture":72,"Logistics":68,"Infrastructure":64,
             "Education":62,"Consumer":60,"Defence":88,"Other":48}
        return m.get(str(row.get("sector","Other")),48)

    def status():
        s = str(row.get("status","")).lower()
        return 100 if "active" in s else 50 if s else 35

    def geography():
        s = str(row.get("state","")).lower()
        top = ["maharashtra","karnataka","delhi","gujarat","telangana","tamil","haryana","pune","mumbai","bengaluru","bangalore"]
        return 100 if any(p in s for p in top) else 65

    def disclosure():
        n = sum(1 for f in ["email","address","activity","roc"] if str(row.get(f,"")).strip())
        return min(30+n*17, 100)

    def signal():
        n = sum(1 for f in ["cin","company_name","state","activity","email","address","company_class"] if str(row.get(f,"")).strip())
        return min(35+n*9, 100)

    def profitability():
        """Revenue potential: high capital + high-margin sectors = higher score"""
        p = float(row.get("paid_up_capital",0) or 0)
        sec = str(row.get("sector","Other"))
        high_margin = {"Technology","Financial Services","Healthcare","Defence","Renewable Energy"}
        base = 70 if sec in high_margin else 50
        if p >= 5e7:  return min(base + 25, 100)
        if p >= 1e7:  return min(base + 15, 100)
        if p >= 5e6:  return min(base + 5,  100)
        return max(base - 10, 20)

    def compliance():
        """MCA compliance: presence of CIN, email, address, status, ROC = strong compliance"""
        fields = ["cin","email","address","status","roc","company_class","activity"]
        n = sum(1 for f in fields if str(row.get(f,"")).strip() not in ("","nan","none","unknown"))
        return min(20 + n * 11, 100)

    def scalability():
        """Scalability based on sector archetype"""
        m = {"Technology":100,"Financial Services":95,"Renewable Energy":92,"Healthcare":88,
             "Defence":85,"Education":72,"Consumer":70,"Logistics":66,"Manufacturing":60,
             "Infrastructure":58,"Agriculture":50,"Other":40}
        return m.get(str(row.get("sector","Other")), 40)

    def promoter():
        """Promoter credibility proxy: private structure + disclosure + vintage"""
        s = 40
        if "private" in str(row.get("company_class","")).lower(): s += 20
        if str(row.get("email","")).strip(): s += 15
        if str(row.get("address","")).strip(): s += 10
        try:
            age = datetime.now().year - int(str(row.get("inc_year","0"))[:4])
            if 3 <= age <= 15: s += 15
        except: pass
        return min(s, 100)

    def ipo_readiness():
        """Composite IPO readiness: capital + vintage + structure + status"""
        score = 0
        p = float(row.get("paid_up_capital",0) or 0)
        if 1e7 <= p <= 2.5e8: score += 30
        elif p > 2.5e8:        score += 20
        elif p >= 5e6:         score += 15
        try:
            age = datetime.now().year - int(str(row.get("inc_year","0"))[:4])
            if 3 <= age <= 12: score += 30
            elif age <= 2:     score += 10
            else:              score += 20
        except: score += 15
        if "private" in str(row.get("company_class","")).lower(): score += 25
        if "active" in str(row.get("status","")).lower():         score += 15
        return min(score, 100)

    def regulatory():
        """Regulatory environment & SEBI appetite by sector"""
        m = {"Renewable Energy":100,"Technology":96,"Financial Services":90,"Defence":88,
             "Healthcare":85,"Infrastructure":80,"Logistics":72,"Manufacturing":68,
             "Education":65,"Consumer":62,"Agriculture":58,"Other":45}
        return m.get(str(row.get("sector","Other")), 45)

    sm = {
        "capital":      cap(),
        "headroom":     headroom(),
        "vintage":      vintage(),
        "structure":    structure(),
        "sector":       sector(),
        "status":       status(),
        "geography":    geography(),
        "disclosure":   disclosure(),
        "signal":       signal(),
        "profitability":profitability(),
        "compliance":   compliance(),
        "scalability":  scalability(),
        "promoter":     promoter(),
        "ipo_readiness":ipo_readiness(),
        "regulatory":   regulatory(),
    }
    tw = sum(w.values())
    return int(round(sum(sm.get(k,0)*w[k] for k in w)/tw))

def apply_scoring(df: pd.DataFrame, W: dict) -> pd.DataFrame:
    df = df.copy()
    df["score"] = df.apply(lambda r: score_one(r,W), axis=1)
    return df.sort_values(["score","paid_up_capital"],ascending=[False,False]).reset_index(drop=True)


# ─── CLAUDE AI ───────────────────────────────────────────────
def get_key() -> str:
    return st.session_state.get("api_key","") or os.getenv("ANTHROPIC_API_KEY","")

def ai_ok() -> bool:
    return bool(get_key())

def call_claude(prompt: str, max_tokens: int=4000) -> str:
    key = get_key()
    if not key:
        return "⚠️ No API key. Add your Anthropic key in the sidebar."
    try:
        res = requests.post(
            ANTHROPIC_URL,
            headers={
                "Content-Type": "application/json",
                "x-api-key": key,
                "anthropic-version": "2023-06-01",
                "anthropic-beta": "web-search-2025-03-05",
            },
            json={
                "model": CLAUDE_MODEL,
                "max_tokens": max_tokens,
                "tools": [{"type":"web_search_20250305","name":"web_search"}],
                "messages": [{"role":"user","content":prompt}],
            },
            timeout=180,
        )
        data = res.json()
        if "error" in data:
            return f"API Error: {data['error'].get('message','Unknown')}"
        texts = [b["text"] for b in data.get("content",[]) if b.get("type")=="text"]
        return "\n".join(texts) or "No response."
    except Exception as e:
        return f"Request failed: {e}"


# ─── DOCUMENT EXPORT ────────────────────────────────────────
def make_doc(title: str, body: str) -> tuple[bytes,str,str]:
    if Document is not None:
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"GreenFlow Ventures  ·  {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph("")
        for blk in body.split("\n\n"):
            doc.add_paragraph(blk.strip())
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(),"application/vnd.openxmlformats-officedocument.wordprocessingml.document","docx"
    html = (f"<html><head><meta charset='utf-8'><title>{title}</title>"
            "<style>body{font-family:Calibri;padding:36px;max-width:900px;margin:0 auto;line-height:1.8}"
            "h1{color:#2D4A1E}pre{white-space:pre-wrap}</style></head>"
            f"<body><h1>{title}</h1><p style='color:#6B8055'>{datetime.now().strftime('%d %B %Y')} · GreenFlow Ventures</p>"
            f"<pre>{body}</pre></body></html>").encode()
    return html,"application/msword","doc"


# ─── STATE ───────────────────────────────────────────────────
def init():
    defs = {
        "W":       {f.key:f.weight for f in FACTORS},
        "df":      pd.DataFrame(),
        "sel":     None,
        "cache":   {},
        "notes":   "",
        "api_key": os.getenv("ANTHROPIC_API_KEY",""),
    }
    for k,v in defs.items():
        if k not in st.session_state:
            st.session_state[k] = v

def get_df() -> pd.DataFrame:
    return st.session_state.get("df", pd.DataFrame())

def get_sel(df: pd.DataFrame) -> pd.Series | None:
    cin = st.session_state.get("sel")
    if not cin or df.empty: return None
    m = df.loc[df["cin"]==cin]
    return m.iloc[0] if not m.empty else None


# ─── SIDEBAR ─────────────────────────────────────────────────
def render_sidebar(base: pd.DataFrame) -> pd.DataFrame:
    with st.sidebar:
        st.markdown("## 🌿 GreenFlow")
        st.caption("Company Intelligence Platform")
        st.divider()

        st.markdown("### Upload MCA File")
        uploaded = st.file_uploader("Excel or CSV", type=["xlsx","xls","csv"], label_visibility="collapsed")
        if uploaded:
            with st.spinner("Parsing file…"):
                parsed = load_file(uploaded)
            if not parsed.empty:
                parsed = apply_scoring(parsed, st.session_state.W)
                st.session_state.df  = parsed
                st.session_state.sel = None
                st.success(f"✓ {len(parsed):,} companies loaded")
                base = parsed.copy()
            else:
                st.error("No valid company records found.")

        if base.empty and not get_df().empty:
            base = get_df().copy()

        st.divider()
        st.markdown("### Claude AI Key")
        k = st.text_input("Key", value=st.session_state.get("api_key",""),
                          type="password", label_visibility="collapsed",
                          placeholder="sk-ant-…")
        if k: st.session_state.api_key = k
        if ai_ok(): st.success("Claude AI active ✓")
        else:        st.warning("Add API key to enable AI analysis")

        st.divider()
        st.markdown("### Scoring Weights")
        st.caption("0 = exclude · 10 = max  |  Default 6.66 × 15 = 100")
        W: dict[str,float] = {}
        for f in FACTORS:
            W[f.key] = float(st.number_input(
                f.label, min_value=0.0, max_value=10.0,
                value=float(st.session_state.W.get(f.key, f.weight)),
                step=0.01, format="%.2f", help=f.desc, key=f"w_{f.key}",
            ))
        st.session_state.W = W
        total_w = round(sum(W.values()), 2)
        st.caption(f"Total weight: **{total_w:.2f}** / 100")
        st.markdown(" ")
        if st.button("▶ Run Scoring", use_container_width=True):
            b = get_df()
            if b.empty: st.warning("Upload a file first.")
            else:
                scored = apply_scoring(b, W)
                st.session_state.df = scored
                st.success(f"Scored {len(scored):,} companies")
                base = scored.copy()

        st.divider()
        st.markdown("### Filter")
        q  = st.text_input("Search", placeholder="Name, CIN, city, sector…", label_visibility="collapsed")
        sa = ["All"]+sorted(base["sector"].dropna().unique()) if not base.empty else ["All"]
        sb = ["All"]+sorted(base["state"].dropna().unique())  if not base.empty else ["All"]
        sc = ["All"]+sorted(base["status"].dropna().unique()) if not base.empty else ["All"]
        fs = st.selectbox("Sector", sa)
        ft = st.selectbox("State",  sb)
        fp = st.selectbox("Status", sc)
        fm = st.number_input("Min Paid-up (₹)", min_value=0.0, value=0.0, step=100000.0)

    if base.empty: return base
    fd = base.copy()
    if q:
        mask = (fd["company_name"].str.contains(q,case=False,na=False)|
                fd["cin"].str.contains(q,case=False,na=False)|
                fd["activity"].str.contains(q,case=False,na=False)|
                fd["state"].str.contains(q,case=False,na=False)|
                fd["email"].str.contains(q,case=False,na=False)|
                fd["address"].str.contains(q,case=False,na=False))
        fd = fd[mask]
    if fs != "All": fd = fd[fd["sector"]==fs]
    if ft != "All": fd = fd[fd["state"]==ft]
    if fp != "All": fd = fd[fd["status"]==fp]
    fd = fd[fd["paid_up_capital"]>=fm]
    return fd.reset_index(drop=True)


# ─── COMPANY PANEL ───────────────────────────────────────────
def show_company(row: pd.Series) -> None:
    sc = int(row["score"])
    c1, c2 = st.columns([3,2])
    with c1:
        st.markdown(f"### {row['company_name']}")
        st.caption(f"{row['cin']}  ·  {row['state']}  ·  {row['company_class']}")
        mc1,mc2,mc3 = st.columns(3)
        mc1.metric("Score",  f"{sc}/100")
        mc2.metric("Rating", score_label(sc))
        mc3.metric("Sector", row["sector"])
        st.progress(sc/100)
    with c2:
        st.markdown("**Key Facts**")
        st.write(f"- **Paid-up:** {inr(float(row['paid_up_capital']))}")
        st.write(f"- **Auth Capital:** {inr(float(row['authorised_capital']))}")
        st.write(f"- **Status:** {row['status']}")
        st.write(f"- **Inc. Year:** {row.get('inc_year','—')}")
        if row.get("email"): st.write(f"- **Email:** {row['email']}")
        if row.get("address"): st.write(f"- **Address:** {str(row['address'])[:80]}")

    lk = links(row["company_name"], row["cin"])
    with st.expander("Research Links"):
        cols = st.columns(len(lk))
        for i,(label,url) in enumerate(lk.items()):
            cols[i].markdown(f"[{label}]({url})")


# ─── TABS ────────────────────────────────────────────────────
def tab_dashboard(df: pd.DataFrame) -> None:
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Companies",   f"{len(df):,}")
    c2.metric("Prime (≥85)", str(int((df["score"]>=85).sum())) if not df.empty else "0")
    c3.metric("Avg Score",   f"{int(df['score'].mean()) if not df.empty else 0}/100")
    c4.metric("Sectors",     str(df["sector"].nunique()) if not df.empty else "0")
    st.divider()

    if df.empty:
        st.info("Upload an MCA Excel or CSV file via the sidebar, then click **Run Scoring**.")
        return

    chosen = st.selectbox(
        "Select a company:",
        options=df["cin"].tolist(),
        format_func=lambda c: f'{df.loc[df["cin"]==c,"company_name"].iloc[0]}  ·  {c}  ·  {df.loc[df["cin"]==c,"state"].iloc[0]}',
    )
    st.session_state.sel = chosen
    row = get_sel(df)
    if row is not None:
        st.divider()
        show_company(row)

    st.divider()
    st.markdown("### Ranked Company List")
    view = df.head(20)[["company_name","cin","sector","state","status","paid_up_capital","score"]].copy()
    view["paid_up_capital"] = view["paid_up_capital"].map(inr)
    view["rating"] = view["score"].map(score_label)
    st.dataframe(view, use_container_width=True, hide_index=True)


def tab_analysis(df: pd.DataFrame) -> None:
    st.markdown("## 360° Deep Dive Analysis")
    st.caption("McKinsey-grade: business model, directors, financials, sector thematic, risk register, recent news, and engagement roadmap.")
    row = get_sel(df)
    if row is None:
        st.info("Select a company from the Dashboard tab first.")
        return
    show_company(row)
    st.divider()

    sf = st.file_uploader("Supporting notes (optional)", type=["txt","csv","xlsx"])
    supp = ""
    if sf:
        try:
            if sf.name.lower().endswith((".xlsx",".xls")): supp = pd.read_excel(sf).head(80).to_csv(index=False)
            elif sf.name.lower().endswith(".csv"):          supp = pd.read_csv(sf).head(80).to_csv(index=False)
            else:                                            supp = sf.read().decode("utf-8","ignore")
        except: pass

    notes = st.text_area("Analyst notes", value=st.session_state.get("notes",""), height=100,
                         placeholder="Known details: founder, contacts, financial data, red flags…")

    if st.button("Run 360° Analysis", use_container_width=True, disabled=not ai_ok()):
        prompt = f"""You are a Principal Analyst at GreenFlow Ventures conducting a comprehensive 360° due diligence analysis for MD review. Search the web extensively: Zauba, Tofler, MCA21, LinkedIn, Google News, BSE SME, Economic Times, Business Standard, IBEF sector reports.

COMPANY: {row["company_name"]} | CIN: {row["cin"]} | State: {row["state"]} | ROC: {row.get("roc","N/A")}
Sector: {row["sector"]} | Activity: {row["activity"]}
Paid-up: {inr(float(row["paid_up_capital"]))} | Auth: {inr(float(row["authorised_capital"]))}
Class: {row["company_class"]} | Status: {row["status"]} | Inc: {row.get("inc_year","—")}
Email: {row.get("email","—")} | Address: {row.get("address","—")}
GreenFlow Score: {int(row["score"])}/100
Notes: {notes or "None."} | Supporting: {supp[:500] if supp else "None."}

PRODUCE ALL 12 SECTIONS — DO NOT SKIP ANY:

1. EXECUTIVE VERDICT
   Three bullets: opportunity, risk, recommended action.

2. BUSINESS INTELLIGENCE
   2.1 Core business model and revenue streams
   2.2 Products / services with market positioning
   2.3 Principal clients / customer segments (name where found)
   2.4 Geographic footprint and delivery model
   2.5 Technology assets and IP

3. FINANCIAL ASSESSMENT
   3.1 Capital structure (paid-up vs authorised utilisation ratio)
   3.2 Revenue/turnover estimate (from Tofler, Zauba, sector multiples)
   3.3 EBITDA margin estimate by sector benchmark
   3.4 Debt profile and leverage indicators
   3.5 Working capital position

4. MANAGEMENT & GOVERNANCE
   4.1 Director names with DIN numbers (from MCA)
   4.2 LinkedIn profiles for key directors / promoters
   4.3 Promoter background — education, prior ventures, track record
   4.4 Board independence and governance quality
   4.5 Any compliance flags or MCA defaults

5. MARKET & SECTOR ANALYSIS
   5.1 Total Addressable Market (₹ Cr) — cite source
   5.2 Serviceable Addressable Market
   5.3 Historical sector CAGR (5 years) — cite source
   5.4 Projected growth (next 5 years) — multiple sources
   5.5 Government policy tailwinds — PLI, budget, regulations
   5.6 PE/VC and institutional activity in this sector
   5.7 Structural demand drivers

6. THEMATIC SECTOR REPORT
   6.1 Latest thematic research from investment banks on this sector
   6.2 Recent IPOs in this sector — listing premium, current performance, P/E
   6.3 Top 5 listed peers — names, market cap, EV/Revenue
   6.4 Highest-growth sub-segments

7. RECENT NEWS & DEVELOPMENTS (last 12 months)
   7.1 Company-specific news and announcements
   7.2 Regulatory or policy changes affecting this sector
   7.3 Competitor moves affecting positioning
   7.4 Negative press or litigation signals

8. COMPETITIVE LANDSCAPE
   8.1 Top 5 direct competitors (name, scale, listed/unlisted)
   8.2 Competitive moat and advantages
   8.3 Market share positioning
   8.4 Porter's Five Forces summary

9. IPO READINESS MATRIX
   9.1 SEBI SME IPO eligibility checklist — each item: ✓ / ✗ / ?
   9.2 Pre-conditions outstanding
   9.3 Restructuring required
   9.4 Estimated months to readiness
   9.5 BSE SME vs NSE Emerge recommendation with rationale

10. RISK REGISTER
    Table: Risk | Category | Severity (H/M/L) | Probability | Mitigation
    Minimum 6 risks: business, financial, regulatory, market, execution, ESG.

11. DEAL THESIS
    11.1 Recommended GreenFlow service (SME IPO / Debt Syndication / Sweat Equity / Pre-IPO)
    11.2 Deal ticket and timeline
    11.3 Valuation benchmarking vs listed peers
    11.4 Value creation roadmap
    11.5 Exit strategy and return framework

12. ENGAGEMENT ROADMAP
    12.1 Personalised founder outreach message — ready to send
    12.2 LinkedIn search query for MD/founder
    12.3 Top 3 value propositions to lead with
    12.4 30-60-90 day action plan for GreenFlow

Write at Goldman Sachs / Kotak IBD standard. Cite sources. Minimum 2,000 words."""

        with st.spinner("Running 360° analysis with live web research (60–90 seconds)…"):
            result = call_claude(prompt, 4000)
        st.session_state.cache[f"deep::{row['cin']}"] = result

    result = st.session_state.cache.get(f"deep::{row['cin']}")
    if result:
        st.divider()
        st.write(result)
        b, mime, ext = make_doc(f"{row['company_name']} — 360° Analysis", result)
        st.download_button("Download Analysis", data=b,
                           file_name=f"{safe_fn(row['company_name'])}-360.{ext}", mime=mime)


def tab_cdr(df: pd.DataFrame) -> None:
    st.markdown("## Company Detailed Report (CDR)")
    st.caption("Formal 12-section investment banking document for board, co-investors, and LP due diligence.")
    row = get_sel(df)
    if row is None:
        st.info("Select a company from the Dashboard tab first.")
        return
    show_company(row)
    st.divider()

    focus = st.text_area("Additional focus areas", height=90,
                         placeholder="E.g. Promoter track record, ESG angle, working capital, IPO timeline…")

    if st.button("Generate CDR", use_container_width=True, disabled=not ai_ok()):
        prompt = f"""You are a Principal at GreenFlow Ventures preparing a Company Detailed Report (CDR) for board presentation, co-investor review, and LP due diligence. Standard: JP Morgan / Kotak pre-IPO research note.

Search: Zauba, Tofler, MCA21, LinkedIn, BSE, NSE, SEBI, IBEF sector reports, Economic Times, Business Standard, Moneycontrol, RBI.

COMPANY: {row["company_name"]} | CIN: {row["cin"]} | {row["state"]}
Sector: {row["sector"]} | Activity: {row["activity"]}
Paid-up: {inr(float(row["paid_up_capital"]))} | Auth: {inr(float(row["authorised_capital"]))}
Class: {row["company_class"]} | Status: {row["status"]} | Inc: {row.get("inc_year","—")}
Score: {int(row["score"])}/100 | Focus: {focus or "Full coverage"} | Context: {st.session_state.get("notes","None.")}

COMPANY DETAILED REPORT
GreenFlow Ventures Advisory · {datetime.now().strftime("%d %B %Y")} · CONFIDENTIAL

SECTION 1 — COMPANY SNAPSHOT: Legal name, CIN, incorporation, ROC, class, status, capitals, utilisation %, age, sector, website, email, address.

SECTION 2 — BUSINESS OVERVIEW: 2.1 Company description (3 paragraphs: founding, evolution, current). 2.2 Business Model Canvas. 2.3 Products & services. 2.4 Geographic footprint. 2.5 Technology & IP.

SECTION 3 — INDUSTRY & MARKET: 3.1 TAM (cite). 3.2 SAM. 3.3 Historical CAGR 5yr (cite). 3.4 Projected growth 5yr (multiple sources). 3.5 PLI & policy. 3.6 Sector IPO activity. 3.7 Demand drivers. 3.8 Investment bank thematic insights. 3.9 Headwinds.

SECTION 4 — FINANCIAL PROFILE: 4.1 Capital structure. 4.2 Revenue estimate. 4.3 Profitability. 4.4 Debt/leverage. 4.5 Working capital. 4.6 Valuation benchmarking (P/E, EV/Revenue vs peers).

SECTION 5 — MANAGEMENT & GOVERNANCE: 5.1 Promoter & founder profile. 5.2 Directors with DIN. 5.3 Board composition. 5.4 Key management. 5.5 MCA compliance.

SECTION 6 — COMPETITIVE ANALYSIS: 6.1 Top 5 competitors. 6.2 Competitive moat. 6.3 Market share. 6.4 Porter's Five Forces.

SECTION 7 — SWOT: Strengths (4), Weaknesses (4), Opportunities (4), Threats (4) with explanations.

SECTION 8 — RISK REGISTER: 8 risks across business/financial/regulatory/market/execution/ESG — each: Description | Severity | Probability | Mitigation.

SECTION 9 — ESG: Environmental, social, governance, ESG IPO positioning.

SECTION 10 — IPO READINESS: SEBI SME checklist (✓/✗/?), pre-conditions, restructuring, timeline, BSE SME vs NSE Emerge.

SECTION 11 — DEAL THESIS: GreenFlow service, deal size, investment thesis, value creation, exit strategy.

SECTION 12 — ENGAGEMENT ROADMAP: Outreach strategy, personalised opening message, LinkedIn query, 30-60-90 day plan.

DISCLAIMER: GreenFlow Ventures internal advisory use only. Not for distribution.

Minimum 2,500 words. Cite all data."""

        with st.spinner("Generating CDR (90 seconds)…"):
            result = call_claude(prompt, 4000)
        st.session_state.cache[f"cdr::{row['cin']}"] = result

    result = st.session_state.cache.get(f"cdr::{row['cin']}")
    if result:
        st.divider()
        st.write(result)
        b, mime, ext = make_doc(f"{row['company_name']} — CDR", result)
        st.download_button("Download CDR", data=b,
                           file_name=f"{safe_fn(row['company_name'])}-cdr.{ext}", mime=mime)


def tab_search(df: pd.DataFrame) -> None:
    st.markdown("## Company Search")
    st.caption("Search by company name, CIN, city, district, state, sector, or email.")
    if df.empty:
        st.info("Upload data first.")
        return
    q = st.text_input("Search", placeholder="e.g. Pune, pharma, U72900MH2018, renewable…")
    res = df.copy()
    if q.strip():
        mask = (res["company_name"].str.contains(q,case=False,na=False)|
                res["cin"].str.contains(q,case=False,na=False)|
                res["activity"].str.contains(q,case=False,na=False)|
                res["state"].str.contains(q,case=False,na=False)|
                res["address"].str.contains(q,case=False,na=False)|
                res["email"].str.contains(q,case=False,na=False))
        res = res[mask]
    st.caption(f"{len(res)} results")
    view = res.head(50)[["company_name","cin","sector","state","status","paid_up_capital","score"]].copy()
    view["paid_up_capital"] = view["paid_up_capital"].map(inr)
    view["rating"] = view["score"].map(score_label)
    st.dataframe(view, use_container_width=True, hide_index=True)
    if not res.empty:
        show_company(res.iloc[0])


def tab_sectors(df: pd.DataFrame) -> None:
    st.markdown("## Sector Discovery")
    if df.empty:
        st.info("Upload data first.")
        return
    secs = sorted(df["sector"].dropna().unique())
    sel = st.selectbox("Select sector", secs)
    lim = int(st.number_input("Companies to show", 5, 200, 20, 5))
    sdf = df[df["sector"]==sel].sort_values(["score","paid_up_capital"],ascending=[False,False]).head(lim)
    st.caption(f"{len(sdf)} companies in **{sel}**")
    view = sdf[["company_name","cin","state","status","paid_up_capital","score"]].copy()
    view["paid_up_capital"] = view["paid_up_capital"].map(inr)
    view["rating"] = view["score"].map(score_label)
    st.dataframe(view, use_container_width=True, hide_index=True)


def tab_top10(df: pd.DataFrame) -> None:
    st.markdown("## Top 10 Intelligence Report")
    if df.empty:
        st.info("Upload data first.")
        return
    top = df.head(10)
    view = top[["company_name","cin","sector","state","paid_up_capital","score"]].copy()
    view["paid_up_capital"] = view["paid_up_capital"].map(inr)
    st.dataframe(view, use_container_width=True, hide_index=True)

    if st.button("Generate Top 10 Report", use_container_width=True, disabled=not ai_ok()):
        listing = "\n".join(
            f'{i+1}. {r["company_name"]} | {r["cin"]} | {r["sector"]} | {r["state"]} | {int(r["score"])}/100 | {inr(float(r["paid_up_capital"]))}'
            for i,r in top.iterrows()
        )
        prompt = f"""Research Director at GreenFlow Ventures. Produce an MD-level intelligence brief on our top 10 SME IPO prospects.

COMPANIES:
{listing}

For each company search Zauba, Tofler, LinkedIn, Google News and provide: director names + LinkedIn, revenue highlights, business model, top 3 competitors, recent news (last 6 months), IPO readiness, recommended GreenFlow service, one priority action.

Open with 200-word executive summary. Close with prioritisation matrix ranking all 10 by GreenFlow revenue potential. Write at Kotak Investment Banking pitch standard."""
        with st.spinner("Generating report (60 seconds)…"):
            result = call_claude(prompt, 4000)
        st.session_state.cache["top10"] = result

    result = st.session_state.cache.get("top10")
    if result:
        st.divider()
        st.write(result)
        b, mime, ext = make_doc("GreenFlow Top 10 IPO Prospects", result)
        st.download_button("Download Report", data=b, file_name=f"greenflow-top10.{ext}", mime=mime)


def tab_sector_research() -> None:
    st.markdown("## Sector Research")
    st.caption("Thematic investment research: market size, CAGR, policy, IPO pipeline, institutional appetite.")
    sec = st.text_input("Sector", placeholder="e.g. Renewable Energy, Fintech, EdTech, Defence…")

    quick = ["Renewable Energy","Pharmaceutical","Technology & SaaS","Fintech & NBFC",
             "Defence","Agritech","EV & Clean Tech","Healthcare","EdTech","Logistics"]
    cols = st.columns(5)
    for i,s in enumerate(quick[:5]):
        if cols[i].button(s, key=f"qa_{i}"):
            sec = s
    cols2 = st.columns(5)
    for i,s in enumerate(quick[5:]):
        if cols2[i].button(s, key=f"qb_{i}"):
            sec = s

    if st.button("Generate Sector Report", use_container_width=True, disabled=not ai_ok()):
        if not sec.strip():
            st.warning("Enter a sector name first.")
            return
        prompt = f"""Professional thematic investment research note on "{sec}" sector in India for GreenFlow Ventures.

Search: IBEF, Kotak, Edelweiss, ICICI Securities, Economic Times, Business Standard, RBI, BSE, SEBI.

Include: market size (₹ Cr and $B with sources), historical CAGR 3 years (cite), projected growth 5 years (multiple sources), government PLI and policy support, major listed players with P/E and valuations, recent SME IPOs and performance, PE/VC activity, structural tailwinds, risks, highest-growth sub-segments, latest thematic reports from investment banks, 3-year outlook.

Write at Emkay / ICICI Securities sector research standard. Cite all data."""
        with st.spinner(f"Researching {sec}…"):
            result = call_claude(prompt, 3000)
        st.session_state.cache[f"sec::{sec}"] = result

    for k,v in st.session_state.cache.items():
        if k.startswith("sec::"):
            sn = k.split("::",1)[1]
            st.divider()
            st.markdown(f"### {sn} — Thematic Report")
            st.write(v)
            b, mime, ext = make_doc(f"{sn} Sector Report", v)
            st.download_button(f"Download {sn} Report", data=b,
                               file_name=f"{safe_fn(sn)}-sector.{ext}", mime=mime, key=f"dl_{sn}")
            break


def tab_context() -> None:
    st.markdown("## Internal Context")
    st.caption("Analyst notes included in all AI analysis and CDR generation.")
    sf = st.file_uploader("Upload internal file (optional)", type=["txt","csv","xlsx"])
    ftxt = ""
    if sf:
        try:
            if sf.name.lower().endswith((".xlsx",".xls")): ftxt = pd.read_excel(sf).head(80).to_csv(index=False)
            elif sf.name.lower().endswith(".csv"):          ftxt = pd.read_csv(sf).head(80).to_csv(index=False)
            else:                                            ftxt = sf.read().decode("utf-8","ignore")
        except: pass
    notes = st.text_area("Notes", value=st.session_state.get("notes",""), height=180,
                         placeholder="Founder notes, diligence inputs, financial details, prior meetings, red flags…")
    if st.button("Save Context", use_container_width=True):
        combined = notes.strip()
        if ftxt: combined = f"{combined}\n\nFile:\n{ftxt}".strip()
        st.session_state.notes = combined
        st.success("Context saved — included in all AI analysis.")
    if st.session_state.get("notes"):
        with st.expander("View saved context"):
            st.text(st.session_state.notes)


# ─── MAIN ────────────────────────────────────────────────────
def main() -> None:
    inject_css()
    init()

    base     = get_df()
    filtered = render_sidebar(base)
    W        = st.session_state.W

    if not filtered.empty:
        filtered = apply_scoring(filtered, W)

    # Header
    st.markdown("# 🌿 GreenFlow Ventures")
    st.caption("Company Level Intelligence Platform  ·  Powered by Claude AI + Live Web Research")
    st.divider()

    # Full scored df for analysis tabs
    full = get_df()
    if not full.empty:
        full = apply_scoring(full, W)

    tabs = st.tabs([
        "Dashboard",
        "360° Analysis",
        "Company CDR",
        "Company Search",
        "Sector Discovery",
        "Top 10 Report",
        "Sector Research",
        "Internal Context",
    ])

    with tabs[0]: tab_dashboard(filtered)
    with tabs[1]: tab_analysis(full if not full.empty else filtered)
    with tabs[2]: tab_cdr(full if not full.empty else filtered)
    with tabs[3]: tab_search(filtered)
    with tabs[4]: tab_sectors(filtered)
    with tabs[5]: tab_top10(filtered)
    with tabs[6]: tab_sector_research()
    with tabs[7]: tab_context()


if __name__ == "__main__":
    main()
