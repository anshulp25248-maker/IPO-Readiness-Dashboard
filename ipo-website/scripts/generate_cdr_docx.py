import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


NAVY = RGBColor(9, 31, 74)
BLACK = RGBColor(17, 24, 39)
RED = RGBColor(190, 18, 60)
GREEN = RGBColor(6, 95, 70)


def clean_text(value):
    text = str(value or "").strip()
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^[-*]\s+", "", text)
    text = re.sub(r"^\d+[\).]\s+", "", text)
    text = text.replace("**", "")
    return text.strip()


def looks_like_heading(line):
    cleaned = clean_text(line)
    if not cleaned:
        return False
    if re.match(r"^(section\s+)?\d+\s*[-.:]", line.strip(), flags=re.I):
        return True
    if line.strip().startswith("#"):
        return True
    if cleaned.endswith(":") and len(cleaned) <= 90:
        return True
    return cleaned.isupper() and len(cleaned.split()) <= 12


def color_for_text(text):
    lower = text.lower()
    if any(term in lower for term in ["red flag", "risk", "litigation", "penalty", "default", "negative", "caution", "discrepancy"]):
        return RED
    if any(term in lower for term in ["positive", "strength", "tailwind", "opportunity", "advantage", "bull case"]):
        return GREEN
    return BLACK


def add_heading(document, text, level=1):
    heading = document.add_heading("", level=level)
    run = heading.add_run(clean_text(text).upper() if level <= 1 else clean_text(text))
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.name = "Arial"
    run.font.size = Pt(15 if level <= 1 else 12.5)
    return heading


def add_body_paragraph(document, text):
    cleaned = clean_text(text)
    if not cleaned:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(cleaned)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = color_for_text(cleaned)


def add_markdown(document, report):
    for raw_line in report.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if looks_like_heading(line):
            add_heading(document, line, level=1 if re.match(r"^(section\s+)?\d+", line, flags=re.I) else 2)
        else:
            add_body_paragraph(document, line)


def main():
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    payload = json.loads(input_path.read_text(encoding="utf-8"))

    company = payload.get("company", {})
    report = payload.get("report", "")
    score = payload.get("score", "NA")
    sources = payload.get("sources", [])

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_heading("", level=0)
    title_run = title.add_run("Scout Smarter Company Diligence Report")
    title_run.bold = True
    title_run.font.color.rgb = NAVY
    title_run.font.name = "Arial"

    for line in [
        f"Company: {company.get('name', 'NA')}",
        f"CIN: {company.get('cin', 'NA')}",
        f"Score: {score}/100",
        f"Paid-up Capital: {company.get('paidUpCapital', 'NA')}",
        f"Authorized Capital: {company.get('authorizedCapital', 'NA')}",
        f"Sector: {company.get('sector', 'NA')}",
        f"Location: {company.get('city', 'NA')}, {company.get('state', 'NA')}",
    ]:
        add_body_paragraph(document, line)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_markdown(document, report)

    if sources:
        add_heading(document, "Source Feed", level=1)
        for source in sources:
            title = source.get("title") or "Source"
            url = source.get("url") or "NA"
            add_body_paragraph(document, f"{title}: {url}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    main()
